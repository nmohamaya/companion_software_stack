#!/usr/bin/env python3
"""
Extract ALL code listings from Process 4-7 docx files.
Reads both paragraphs and tables, captures file paths and code blocks.
"""

import os
from docx import Document

BASE_DIR = "/home/nmohanan31/NM/Projects/Perception"

DOCS = [
    "Process_4_Mission_Planner.docx",
    "Process_5_Comms.docx",
    "Process_6_Payload_Manager.docx",
    "Process_7_System_Monitor.docx",
]


def extract_from_doc(filepath):
    """Extract all text from paragraphs and tables in order."""
    doc = Document(filepath)
    all_text = []

    # We need to iterate in document order (paragraphs and tables interleaved)
    # python-docx body contains both paragraphs and tables in order
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            # It's a paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is not None:
                style_name = para.style.name if para.style else ""
                all_text.append(("PARA", para.text, style_name))
        elif tag == 'tbl':
            # It's a table
            tbl = None
            for t in doc.tables:
                if t._element is element:
                    tbl = t
                    break
            if tbl is not None:
                for row_idx, row in enumerate(tbl.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text
                        row_cells.append(cell_text)
                    all_text.append(("TABLE_ROW", row_cells, f"row_{row_idx}"))
    return all_text


def process_document(doc_path):
    """Process a single document and write full output."""
    doc_name = os.path.basename(doc_path).replace('.docx', '')
    print(f"\n{'='*120}")
    print(f"DOCUMENT: {doc_name}")
    print(f"{'='*120}")

    elements = extract_from_doc(doc_path)

    for i, (etype, content, style) in enumerate(elements):
        if etype == "PARA":
            if content.strip():
                print(f"[PARA|{style}] {content}")
        elif etype == "TABLE_ROW":
            cells_str = " | ".join(str(c).replace('\n', '\\n') for c in content)
            print(f"[TABLE_ROW] {cells_str}")


def main():
    for doc_name in DOCS:
        doc_path = os.path.join(BASE_DIR, doc_name)
        if os.path.exists(doc_path):
            process_document(doc_path)
        else:
            print(f"WARNING: {doc_path} not found!")


if __name__ == "__main__":
    main()
