#!/usr/bin/env python3
"""
Extract ALL code listings from Process 4-7 docx files.
Reads both paragraphs and tables, saves full output per document.
"""

import os
from docx import Document

BASE_DIR = "/home/nmohanan31/NM/Projects/Perception"
OUT_DIR  = os.path.join(BASE_DIR, "extracted_code")

DOCS = [
    "Process_4_Mission_Planner.docx",
    "Process_5_Comms.docx",
    "Process_6_Payload_Manager.docx",
    "Process_7_System_Monitor.docx",
]


def extract_from_doc(filepath):
    """Extract all text from paragraphs and tables in document order."""
    doc = Document(filepath)
    all_elements = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            for p in doc.paragraphs:
                if p._element is element:
                    style_name = p.style.name if p.style else ""
                    all_elements.append(("PARA", p.text, style_name))
                    break
        elif tag == 'tbl':
            for t in doc.tables:
                if t._element is element:
                    for row_idx, row in enumerate(t.rows):
                        row_cells = []
                        for cell in row.cells:
                            row_cells.append(cell.text)
                        all_elements.append(("TABLE_ROW", row_cells, f"row_{row_idx}"))
                    break
    return all_elements


def is_code_content(text):
    """Check if text looks like code."""
    code_indicators = [
        '#include', '#pragma', 'class ', 'struct ', 'void ', 'int ', 'double ',
        'namespace ', 'template', 'public:', 'private:', 'protected:',
        'return ', 'if (', 'while (', 'for (', 'auto ', 'const ',
        '// ', '/*', '};', '{}', 'std::', 'cmake_minimum',
        'find_package', 'target_link', 'add_executable', 'add_library',
        'enum class', 'virtual ', 'override', 'explicit ',
        'spdlog::', 'Eigen::', 'mmap(', 'shm_open',
        'atomic<', 'mutex', 'thread', 'static_cast',
        '"name":', '"version":', '"type":', '"waypoints":', # JSON
    ]
    return any(ind in text for ind in code_indicators)


def extract_code_blocks(elements):
    """
    Identify code blocks from the element stream.
    A code block is typically a single table cell containing code,
    preceded by a paragraph naming it (e.g., "Listing X-Y: ...").
    """
    blocks = []
    current_label = ""
    block_num = 0
    
    for i, (etype, content, style) in enumerate(elements):
        if etype == "PARA":
            text = content.strip()
            if text:
                current_label = text
                
        elif etype == "TABLE_ROW":
            # Table rows - each cell might contain code
            for cell_text in content:
                if is_code_content(cell_text) and len(cell_text) > 50:
                    block_num += 1
                    blocks.append({
                        'num': block_num,
                        'label': current_label,
                        'code': cell_text,
                    })
                elif len(cell_text) > 100 and ('💡' in cell_text or 'Concept' in cell_text):
                    # C++ concept explanations - include for context
                    block_num += 1
                    blocks.append({
                        'num': block_num,
                        'label': "CONCEPT: " + current_label,
                        'code': cell_text,
                    })
    
    return blocks


def process_document(doc_name):
    """Process a single document."""
    doc_path = os.path.join(BASE_DIR, doc_name)
    short_name = doc_name.replace('.docx', '')
    out_dir = os.path.join(OUT_DIR, short_name)
    os.makedirs(out_dir, exist_ok=True)
    
    print(f"\n{'='*80}")
    print(f"Processing: {doc_name}")
    print(f"{'='*80}")
    
    elements = extract_from_doc(doc_path)
    
    # Save full raw output
    full_path = os.path.join(out_dir, "full_output.txt")
    with open(full_path, 'w') as f:
        for etype, content, style in elements:
            if etype == "PARA":
                if content.strip():
                    f.write(f"[PARA|{style}] {content}\n")
            elif etype == "TABLE_ROW":
                for ci, cell in enumerate(content):
                    f.write(f"[TABLE_CELL_{ci}] {cell}\n")
                f.write("[END_TABLE_ROW]\n")
    
    # Extract code blocks
    blocks = extract_code_blocks(elements)
    
    for block in blocks:
        block_path = os.path.join(out_dir, f"block_{block['num']:03d}.txt")
        with open(block_path, 'w') as f:
            f.write(f"=== {block['label']} ===\n\n")
            f.write(block['code'])
            f.write("\n")
    
    print(f"  Elements: {len(elements)}")
    print(f"  Code blocks: {len(blocks)}")
    for b in blocks:
        print(f"    Block {b['num']:03d}: {b['label'][:80]}")
    
    return blocks


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    
    all_blocks = {}
    for doc_name in DOCS:
        doc_path = os.path.join(BASE_DIR, doc_name)
        if os.path.exists(doc_path):
            blocks = process_document(doc_name)
            all_blocks[doc_name] = blocks
        else:
            print(f"WARNING: {doc_path} not found!")
    
    print(f"\n{'='*80}")
    print("EXTRACTION COMPLETE")
    print(f"{'='*80}")
    for doc_name, blocks in all_blocks.items():
        print(f"  {doc_name}: {len(blocks)} blocks")


if __name__ == "__main__":
    main()
